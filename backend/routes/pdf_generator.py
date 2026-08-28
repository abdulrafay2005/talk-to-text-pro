"""
pdf_generator.py
Generates and downloads a PDF and a Word (.docx) report for a meeting.
The PDF uses ReportLab, the Word export uses python-docx.
Both contain the same useful sections:
Title, date, summary, key points, decisions, action items, speakers,
and the transcript.
"""

import os

from bson.objectid import ObjectId
from flask import Blueprint, jsonify, send_file
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from xml.sax.saxutils import escape as escape_xml

from database.connection import get_meetings_collection
from models.meeting import convert_meeting_to_view, string_to_object_id
from utils.auth import get_current_user_id, login_required

pdf_bp = Blueprint("pdf", __name__)

# The generated exports folder lives inside the backend folder.
GENERATED_PDFS_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), "generated_pdfs")


def safe(text):
    """
    Escape XML special characters so they can be used inside
    ReportLab Paragraphs without breaking them.
    """
    if text is None:
        return ""
    return escape_xml(str(text))


def get_meeting(meeting_id, user_id):
    """
    Fetch a meeting from MongoDB by id, but only if it belongs to the user.
    Returns None if not found.
    """
    object_id = string_to_object_id(meeting_id)
    if object_id is None:
        return None

    meeting = get_meetings_collection().find_one(
        {"_id": object_id, "user_id": user_id}
    )
    if meeting is None:
        return None

    return convert_meeting_to_view(meeting)


def _safe_download_name(title):
    """
    Build a short, filesystem-safe download name from the meeting title.
    """
    name = (title or "meeting").strip()[:50]
    name = name.replace("/", "-").replace("\\", "-").replace("\x00", "-")
    return name or "meeting"


def generate_meeting_pdf(meeting):
    """
    Create the PDF file for a meeting and return its file path.
    """
    os.makedirs(GENERATED_PDFS_FOLDER, exist_ok=True)

    filename = f"meeting_{meeting['_id']}.pdf"
    file_path = os.path.join(GENERATED_PDFS_FOLDER, filename)

    doc = SimpleDocTemplate(file_path, pagesize=A4)
    styles = getSampleStyleSheet()

    story = []

    # Header
    story.append(Paragraph("Talk To Text Pro - Meeting Report", styles["Title"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph(f"<b>Title:</b> {safe(meeting['title'])}", styles["Normal"]))
    story.append(Paragraph(f"<b>Date:</b> {safe(meeting.get('created_at', ''))}", styles["Normal"]))
    story.append(Paragraph(f"<b>Meeting type:</b> {safe(meeting['meeting_type'])}", styles["Normal"]))
    story.append(Paragraph(f"<b>Original language:</b> {safe(meeting['original_language'])}", styles["Normal"]))
    story.append(Paragraph(f"<b>Duration:</b> {safe(meeting['duration'])} seconds", styles["Normal"]))
    story.append(Paragraph(f"<b>Sentiment:</b> {safe(meeting['sentiment'])}", styles["Normal"]))

    # Summary
    story.append(Spacer(1, 14))
    story.append(Paragraph("Summary", styles["Heading2"]))
    story.append(Paragraph(safe(meeting["summary"]), styles["BodyText"]))

    # Simple helper to add bullet lists
    def add_bullets(title, items):
        story.append(Spacer(1, 10))
        story.append(Paragraph(title, styles["Heading2"]))
        if not items:
            story.append(Paragraph("Nothing was explicitly stated.", styles["BodyText"]))
        else:
            for item in items:
                story.append(Paragraph(f"- {safe(item)}", styles["BodyText"]))

    add_bullets("Key Points", meeting["key_points"])
    add_bullets("Topics", meeting["topics"])
    add_bullets("Decisions", meeting["decisions"])
    add_bullets("Unresolved Issues", meeting["unresolved_issues"])

    # Action items as a table
    story.append(Spacer(1, 10))
    story.append(Paragraph("Action Items", styles["Heading2"]))

    table_data = [["Person", "Task", "Deadline"]]
    for item in meeting["action_items"]:
        table_data.append([safe(item["person"]), safe(item["task"]), safe(item["deadline"])])

    if len(table_data) > 1:
        table = Table(table_data)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2563eb")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("PADDING", (0, 0), (-1, -1), 6),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
        ]))
        story.append(table)
    else:
        story.append(Paragraph("No action items were explicitly stated.", styles["BodyText"]))

    # Transcripts
    story.append(Spacer(1, 14))
    story.append(Paragraph("Transcripts", styles["Heading1"]))

    story.append(Paragraph("Cleaned Transcript", styles["Heading2"]))
    story.append(Paragraph(safe(meeting["transcript_cleaned"]), styles["BodyText"]))

    story.append(Spacer(1, 10))
    story.append(Paragraph("Raw Transcript", styles["Heading2"]))
    story.append(Paragraph(safe(meeting["transcript_raw"]), styles["BodyText"]))

    if meeting.get("translated_transcript"):
        story.append(Spacer(1, 10))
        story.append(Paragraph(f"Translated Transcript ({safe(meeting['translate_to'])})", styles["Heading2"]))
        story.append(Paragraph(safe(meeting["translated_transcript"]), styles["BodyText"]))

    doc.build(story)
    return file_path


@pdf_bp.route("/api/meetings/<meeting_id>/pdf", methods=["GET"])
@login_required
def download_pdf(meeting_id):
    """
    Generate and download the PDF report for a meeting (ownership checked).
    """
    meeting = get_meeting(meeting_id, get_current_user_id())
    if meeting is None:
        return jsonify({"error": "Meeting not found."}), 404

    file_path = generate_meeting_pdf(meeting)

    # Build a safe download filename from the meeting title.
    download_name = f"{_safe_download_name(meeting['title'])}.pdf"

    return send_file(
        file_path,
        as_attachment=True,
        download_name=download_name,
        mimetype="application/pdf",
    )


def generate_meeting_docx(meeting):
    """
    Create the Word (.docx) file for a meeting and return its file path.
    Uses python-docx. Contains the same sections as the PDF export.
    """
    from docx import Document

    os.makedirs(GENERATED_PDFS_FOLDER, exist_ok=True)

    filename = f"meeting_{meeting['_id']}.docx"
    file_path = os.path.join(GENERATED_PDFS_FOLDER, filename)

    document = Document()

    # Header
    document.add_heading("Talk To Text Pro - Meeting Report", level=0)
    document.add_heading("Title: " + str(meeting.get("title", "")), level=1)
    document.add_paragraph("Date: " + str(meeting.get("created_at", "")))
    document.add_paragraph("Meeting type: " + str(meeting.get("meeting_type", "")))
    document.add_paragraph("Original language: " + str(meeting.get("original_language", "")))
    document.add_paragraph("Duration: " + str(meeting.get("duration", "")) + " seconds")
    document.add_paragraph("Sentiment: " + str(meeting.get("sentiment", "")))

    def add_section(title, items):
        document.add_heading(title, level=2)
        if not items:
            document.add_paragraph("Nothing was explicitly stated.")
        else:
            for item in items:
                document.add_paragraph(str(item), style="List Bullet")

    # Summary
    document.add_heading("Summary", level=2)
    document.add_paragraph(str(meeting.get("summary", "")))

    add_section("Key Points", meeting["key_points"])
    add_section("Topics", meeting["topics"])
    add_section("Decisions", meeting["decisions"])
    add_section("Unresolved Issues", meeting["unresolved_issues"])

    # Action items as a table
    document.add_heading("Action Items", level=2)
    action_items = meeting["action_items"] or []
    if action_items:
        table = document.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Person", "Task", "Deadline"
        for item in action_items:
            cells = table.add_row().cells
            cells[0].text = str(item.get("person", ""))
            cells[1].text = str(item.get("task", ""))
            cells[2].text = str(item.get("deadline", ""))
    else:
        document.add_paragraph("No action items were explicitly stated.")

    # Transcripts
    document.add_heading("Transcripts", level=1)
    document.add_heading("Cleaned Transcript", level=2)
    document.add_paragraph(str(meeting["transcript_cleaned"]))

    document.add_heading("Raw Transcript", level=2)
    document.add_paragraph(str(meeting["transcript_raw"]))

    if meeting.get("translated_transcript"):
        document.add_heading(
            "Translated Transcript ({})".format(meeting["translate_to"]),
            level=2,
        )
        document.add_paragraph(str(meeting["translated_transcript"]))

    document.save(file_path)
    return file_path


@pdf_bp.route("/api/meetings/<meeting_id>/docx", methods=["GET"])
@login_required
def download_docx(meeting_id):
    """
    Generate and download the Word (.docx) report for a meeting
    (ownership checked).
    """
    meeting = get_meeting(meeting_id, get_current_user_id())
    if meeting is None:
        return jsonify({"error": "Meeting not found."}), 404

    file_path = generate_meeting_docx(meeting)

    download_name = f"{_safe_download_name(meeting['title'])}.docx"

    return send_file(
        file_path,
        as_attachment=True,
        download_name=download_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )