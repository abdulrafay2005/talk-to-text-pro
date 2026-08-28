"""
test_exports.py
Verifies the PDF and Word (.docx) exports actually produce real, readable
files (not fake output):

    - PDF begins with the %PDF header and has a meaningful size
    - DOCX opens as a real Word package (python-docx can re-open it)
    - both contain the key sections (title, decisions, action items,
      transcript)

Run with:  python -m unittest test_exports -v
"""

import base64
import os
import re
import sys
import unittest
import zlib

# Make the backend folder importable.
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from bson.objectid import ObjectId

from routes.pdf_generator import generate_meeting_docx, generate_meeting_pdf


def pdf_plain_text(path):
    """
    Extract the (possibly ASCII85 + FlateDecode-compressed) text content of a
    PDF so the test can verify the report's sections really are inside the file.
    """
    with open(path, "rb") as handle:
        data = handle.read()

    chunks = []
    for match in re.finditer(rb"stream\r?\n(.*?)endstream", data, re.DOTALL):
        raw = match.group(1)
        # ReportLab writes streams as ASCII85 then Flate.
        try:
            decoded = base64.a85decode(raw, adobe=True)
        except ValueError:
            decoded = raw
        try:
            chunks.append(zlib.decompress(decoded))
        except zlib.error:
            chunks.append(decoded)
    return b"\n".join(chunks)


def make_meeting():
    """A complete meeting dictionary exactly like the real view shape."""
    return {
        "_id": str(ObjectId()),
        "title": "Quarterly Planning",
        "created_at": "2026-08-28 10:00:00",
        "meeting_type": "Planning Meeting",
        "original_language": "en",
        "duration": 1460.0,
        "sentiment": "Positive",
        "meeting_type_value": "Planning Meeting",
        "summary": "We aligned on Q4 targets.",
        "key_points": ["Budget approved", "Remote by default"],
        "topics": ["budget", "hiring"],
        "decisions": ["Budget increased 10%"],
        "unresolved_issues": ["Office lease"],
        "action_items": [
            {"person": "Ali", "task": "ship the API", "deadline": "Friday"},
            {"person": "Sarah", "task": "draft the proposal", "deadline": "Monday"},
        ],
        "transcript_cleaned": (
            "Speaker 1: The budget deadline is March 15th. "
            "Speaker 2: We decided to hire two engineers."
        ),
        "transcript_raw": (
            "um Speaker 1: The budget deadline is March 15th. "
            "Speaker 2: We decided to hire two engineers."
        ),
        "translated_transcript": None,
        "translate_to": None,
        "optimization": {
            "original_words": 18,
            "optimized_words": 18,
            "reduction_percentage": 0.0,
        },
    }


class PdfExportTests(unittest.TestCase):

    def test_pdf_is_generated_and_is_a_real_pdf(self):
        meeting = make_meeting()
        path = generate_meeting_pdf(meeting)

        self.assertTrue(os.path.exists(path))
        self.assertGreater(os.path.getsize(path), 1000)

        with open(path, "rb") as handle:
            header = handle.read(5)
        self.assertEqual(header, b"%PDF-", "Generated file is not a PDF")

        os.remove(path)

    def test_pdf_key_sections_are_embedded(self):
        meeting = make_meeting()
        path = generate_meeting_pdf(meeting)

        content = pdf_plain_text(path)

        # The sections must actually be inside the PDF (after decompression).
        self.assertIn(b"Quarterly Planning", content)
        self.assertIn(b"ship the API", content)
        self.assertIn(b"Budget increased 10%", content)
        self.assertIn(b"Friday", content)
        self.assertIn(
            b"Speaker 1: The budget deadline is March 15th.",
            content,
        )

        os.remove(path)


class DocxExportTests(unittest.TestCase):

    def test_docx_is_generated_and_reopens_as_word_document(self):
        from docx import Document

        meeting = make_meeting()
        path = generate_meeting_docx(meeting)

        self.assertTrue(os.path.exists(path))
        self.assertGreater(os.path.getsize(path), 2000)

        # python-docx re-opening proves it is a valid Word file.
        document = Document(path)

        full_text = "\n".join(
            paragraph.text for paragraph in document.paragraphs
        )
        # Action items live in the table, so include every cell's text too.
        full_text += "\n" + "\n".join(
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        )

        self.assertIn("Quarterly Planning", full_text)
        self.assertIn("Budget increased 10%", full_text)
        self.assertIn("ship the API", full_text)
        self.assertIn(
            "Speaker 1: The budget deadline is March 15th.",
            full_text,
        )

        # Action item table exists with header row "Person / Task / Deadline".
        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual(
            [cell.text for cell in table.rows[0].cells],
            ["Person", "Task", "Deadline"],
        )
        self.assertEqual(table.rows[1].cells[2].text, "Friday")

        os.remove(path)


if __name__ == "__main__":
    unittest.main()